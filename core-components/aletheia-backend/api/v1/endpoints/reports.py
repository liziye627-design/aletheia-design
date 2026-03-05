"""
报告API端点
"""

import base64
import json
from io import BytesIO
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException, Query, status
from pydantic import BaseModel, Field

from core.sqlite_database import get_sqlite_db
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

router = APIRouter()


class GenerateReportRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    content: str = Field(..., min_length=1, max_length=50000)
    template_id: Optional[str] = None
    credibility_score: Optional[float] = Field(default=0.5, ge=0.0, le=1.0)
    tags: List[str] = Field(default_factory=list)
    sources: List[Dict[str, Any]] = Field(default_factory=list)


class ReportResponse(BaseModel):
    id: str
    title: str
    summary: str
    content_html: str
    credibility_score: float
    status: str = "complete"
    created_at: str
    updated_at: str
    sources: List[Dict[str, Any]]
    tags: List[str]


class ReportListResponse(BaseModel):
    items: List[ReportResponse]
    total: int
    page: int
    page_size: int
    has_more: bool


class ReportExportRequest(BaseModel):
    format: str = Field(..., pattern="^(pdf|docx|md|txt|html|json)$")
    payload: Dict[str, Any]


class ReportExportResponse(BaseModel):
    file_name: str
    mime_type: str
    content_base64: str


class GenerateReportFromRunRequest(BaseModel):
    run_id: str = Field(..., min_length=1)
    title: Optional[str] = Field(default=None, max_length=200)
    tags: List[str] = Field(default_factory=list)


def _to_html(content: str) -> str:
    escaped = (
        content.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return f"<article><p>{escaped}</p></article>"


def _build_summary(content: str, max_len: int = 120) -> str:
    text = content.strip().replace("\n", " ")
    if len(text) <= max_len:
        return text
    return text[: max_len - 1] + "…"


def _to_pdf_bytes(title: str, content: str) -> bytes:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    c.setFont("STSong-Light", 16)
    c.drawString(40, height - 50, title)

    c.setFont("STSong-Light", 11)
    y = height - 80
    for raw_line in content.splitlines() or [content]:
        line = raw_line if raw_line.strip() else " "
        while len(line) > 50:
            c.drawString(40, y, line[:50])
            line = line[50:]
            y -= 18
            if y < 40:
                c.showPage()
                c.setFont("STSong-Light", 11)
                y = height - 50
        c.drawString(40, y, line)
        y -= 18
        if y < 40:
            c.showPage()
            c.setFont("STSong-Light", 11)
            y = height - 50

    c.save()
    return buffer.getvalue()


def _to_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.splitlines() or [content]:
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.get("/", response_model=ReportListResponse)
async def get_reports(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
):
    """获取报告列表"""
    try:
        db = get_sqlite_db()
        offset = (page - 1) * page_size
        items_raw = db.list_reports(limit=page_size, offset=offset)
        total = db.count_reports()

        items = [ReportResponse(**item) for item in items_raw]
        has_more = total > page * page_size

        return ReportListResponse(
            items=items,
            total=total,
            page=page,
            page_size=page_size,
            has_more=has_more,
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Get reports failed: {str(e)}",
        )


@router.get("/{report_id}", response_model=ReportResponse)
async def get_report(report_id: str):
    """获取报告详情"""
    db = get_sqlite_db()
    item = db.get_report(report_id)
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Report not found"
        )
    return ReportResponse(**item)


@router.post("/generate", response_model=ReportResponse)
async def generate_report(request: GenerateReportRequest):
    """生成新报告并入库"""
    try:
        db = get_sqlite_db()
        now = datetime.utcnow().isoformat()
        report_id = f"report_{int(datetime.utcnow().timestamp() * 1000)}"

        item = {
            "id": report_id,
            "title": request.title,
            "summary": _build_summary(request.content),
            "content_html": _to_html(request.content),
            "credibility_score": request.credibility_score or 0.5,
            "status": "complete",
            "created_at": now,
            "updated_at": now,
            "sources": request.sources,
            "tags": request.tags,
        }

        db.save_report(item)
        saved = db.get_report(report_id)
        if not saved:
            raise RuntimeError("Save report failed")

        return ReportResponse(**saved)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Generate report failed: {str(e)}",
        )


@router.post("/generate-from-run", response_model=ReportResponse)
async def generate_report_from_run(request: GenerateReportFromRunRequest):
    """从 investigation run 生成报告并入库。"""
    from services.investigation_engine import get_investigation_manager

    run = get_investigation_manager().get_run(request.run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Investigation run not found")
    result = run.get("result")
    if not result:
        raise HTTPException(status_code=400, detail="Investigation is not completed yet")

    keyword = (
        (((result.get("request") or {}).get("keyword")) or result.get("search", {}).get("keyword") or "核验报告")
        if isinstance(result, dict)
        else "核验报告"
    )
    title = request.title or f"Aletheia核验报告-{keyword}-{datetime.utcnow().date()}"
    sections = result.get("report_sections") or []
    section_text = []
    for sec in sections:
        section_text.append(f"## {sec.get('title', '章节')}")
        section_text.append(str(sec.get("content_markdown", "")))
        section_text.append("")
    body = "\n".join(section_text).strip()
    if not body:
        body = json.dumps(result, ensure_ascii=False, indent=2)

    status_value = str(result.get("status") or "complete")
    credibility_score = float(
        (((result.get("dual_profile_result") or {}).get("combined_result") or {}).get("score"))
        or (((result.get("enhanced") or {}).get("reasoning_chain") or {}).get("final_score"))
        or 0.5
    )
    sources = result.get("external_sources") or []

    db = get_sqlite_db()
    now = datetime.utcnow().isoformat()
    report_id = f"report_{int(datetime.utcnow().timestamp() * 1000)}"
    item = {
        "id": report_id,
        "title": title,
        "summary": _build_summary(body),
        "content_html": _to_html(body),
        "credibility_score": credibility_score,
        "status": status_value,
        "created_at": now,
        "updated_at": now,
        "sources": sources,
        "tags": list(set(["from_run", f"run_{request.run_id}"] + request.tags)),
    }
    db.save_report(item)
    saved = db.get_report(report_id)
    if not saved:
        raise HTTPException(status_code=500, detail="Save report failed")
    return ReportResponse(**saved)


@router.post("/export", response_model=ReportExportResponse)
async def export_report(request: ReportExportRequest):
    """导出报告内容（返回base64编码文件内容）"""
    payload = request.payload
    title = str(payload.get("title") or "report")
    content = str(payload.get("content") or "")

    # 基于当前依赖，先提供可稳定使用的文本类导出。
    # pdf/docx在无第三方库时以文本容器降级，确保链路可用。
    if request.format == "md":
        file_name = f"{title}.md"
        mime_type = "text/markdown; charset=utf-8"
        raw = content.encode("utf-8")
    elif request.format == "txt":
        file_name = f"{title}.txt"
        mime_type = "text/plain; charset=utf-8"
        raw = content.encode("utf-8")
    elif request.format == "html":
        file_name = f"{title}.html"
        mime_type = "text/html; charset=utf-8"
        raw = _to_html(content).encode("utf-8")
    elif request.format == "json":
        file_name = f"{title}.json"
        mime_type = "application/json; charset=utf-8"
        raw = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    elif request.format == "pdf":
        file_name = f"{title}.pdf"
        mime_type = "application/pdf"
        raw = _to_pdf_bytes(title, content)
    else:  # docx
        file_name = f"{title}.docx"
        mime_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        raw = _to_docx_bytes(title, content)

    return ReportExportResponse(
        file_name=file_name,
        mime_type=mime_type,
        content_base64=base64.b64encode(raw).decode("ascii"),
    )
